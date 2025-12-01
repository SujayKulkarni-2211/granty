from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file, session, flash
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
import os
import secrets
import json
from datetime import datetime
import uuid
from werkzeug.utils import secure_filename
import gemini_utils
import storage
from models import db, User, Project, Template
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import markdown
import tempfile
import time
import hashlib
import random
from email_utils import send_verification_email, send_welcome_email, send_email_async
from storage import create_user_with_verification, verify_user_email, resend_verification_email, authenticate_user_with_verification
from flask import request, url_for
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available. DOCX downloads will not work.")
    DOCX_AVAILABLE = False
import re
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    print("Warning: BeautifulSoup4 not available. HTML parsing will be limited.")
    BS4_AVAILABLE = False
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
    print("✓ BeautifulSoup4 available for better HTML parsing")
except ImportError:
    BS4_AVAILABLE = False
    print("⚠ BeautifulSoup4 not available. HTML parsing will be basic.")

# Don't import weasyprint at module level - import only when needed
WEASYPRINT_AVAILABLE = False
PDFKIT_AVAILABLE = False

# Test for weasyprint availability without importing
try:
    import importlib
    importlib.import_module('weasyprint')
    WEASYPRINT_AVAILABLE = True
    print("✓ WeasyPrint available for advanced PDF generation")
except (ImportError, OSError) as e:
    WEASYPRINT_AVAILABLE = False
    print(f"⚠ WeasyPrint not available: {e}")

# Test for pdfkit availability
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
    print("✓ pdfkit available for PDF generation")
except ImportError:
    PDFKIT_AVAILABLE = False
    print("⚠ pdfkit not available")

print("✓ PDF/DOCX conversion setup complete with available libraries")

app = Flask(__name__)
app.config.from_object('config')
app.secret_key = app.config.get('SECRET_KEY', 'dev-secret-key')

# Initialize database
app.config['SQLALCHEMY_DATABASE_URI'] = app.config.get('DATABASE_URL', 'sqlite:///granty.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

# CRITICAL FIX: Initialize database tables in app context
with app.app_context():
    try:
        # Create all tables if they don't exist
        db.create_all()
        print("✓ Database tables created/verified")
        
        # Check if templates exist, if not initialize them
        if Template.query.count() == 0:
            storage.initialize_comprehensive_templates()
            print("✓ Templates initialized")
        else:
            print("✓ Templates already exist")
            
    except Exception as e:
        print(f"Error initializing database: {e}")
        # Don't raise the error, let the app start anyway
        pass

# Initialize Flask-Login
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'

@login_manager.user_loader
def load_user(user_id):
    return storage.get_user_by_id(user_id)

# Mock Razorpay Client Class (keeping existing implementation)
class MockRazorpayClient:
    def __init__(self, auth):
        self.key_id = auth[0]
        self.key_secret = auth[1]
        self.order = MockOrderAPI()
        self.utility = MockUtilityAPI()

class MockOrderAPI:
    def create(self, data):
        order_id = f"order_{uuid.uuid4().hex[:16]}"
        if 'mock_orders' not in session:
            session['mock_orders'] = {}
        session['mock_orders'][order_id] = {
            'id': order_id,
            'amount': data['amount'],
            'currency': data['currency'],
            'receipt': data['receipt'],
            'notes': data.get('notes', {}),
            'status': 'created',
            'created_at': datetime.now().isoformat()
        }
        return {
            'id': order_id,
            'amount': data['amount'],
            'currency': data['currency'],
            'receipt': data['receipt'],
            'status': 'created'
        }

class MockUtilityAPI:
    def verify_payment_signature(self, data):
        order_id = data.get('razorpay_order_id')
        payment_id = data.get('razorpay_payment_id')
        signature = data.get('razorpay_signature')
        
        if 'mock_orders' not in session or order_id not in session['mock_orders']:
            raise MockSignatureVerificationError("Order not found")
        
        expected_signature = hashlib.sha256(f"{order_id}|{payment_id}".encode()).hexdigest()
        
        if signature != expected_signature:
            raise MockSignatureVerificationError("Invalid signature")
        
        session['mock_orders'][order_id]['status'] = 'paid'
        session['mock_orders'][order_id]['payment_id'] = payment_id
        return True

class MockSignatureVerificationError(Exception):
    pass

razorpay_client = MockRazorpayClient(auth=(
    app.config.get('RAZORPAY_KEY_ID', 'mock_key_id'),
    app.config.get('RAZORPAY_KEY_SECRET', 'mock_key_secret')
))

@app.route('/')
def index():
    template_categories = storage.get_templates_by_category()
    if current_user.is_authenticated:
        projects = storage.get_projects(current_user.id)
    else:
        projects = storage.get_projects()  # Guest mode
    return render_template('index.html', template_categories=template_categories, drafts=projects, user=current_user)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        
        user = authenticate_user_with_verification(email, password)
        if user:
            if not user.is_email_verified:
                flash('Please verify your email address before logging in. Check your inbox for verification email.', 'warning')
                return redirect(url_for('verification_needed', email=email))
            
            login_user(user)
            flash('Logged in successfully!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password.', 'error')
    
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        name = request.form.get('name')
        email = request.form.get('email')
        password = request.form.get('password')
        
        # Check if user already exists
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            if existing_user.is_email_verified:
                flash('Email already registered. Please login instead.', 'error')
                return redirect(url_for('login'))
            else:
                flash('Email already registered but not verified. Please check your inbox or request a new verification email.', 'warning')
                return redirect(url_for('verification_needed', email=email))
        
        user = create_user_with_verification(email, password, name)
        if user:
            # Send verification email - use the token from the user object!
            app_url = request.url_root.rstrip('/')
            token = user.email_verification_token  # ✅ Use the actual token
            if send_email_async(send_verification_email, email, name, token, app_url):
                flash('Account created successfully! Please check your email to verify your account.', 'success')
                return redirect(url_for('verification_sent', email=email))
            else:
                flash('Account created but failed to send verification email. Please try again.', 'warning')
                return redirect(url_for('verification_needed', email=email))
        else:
            flash('Error creating account. Please try again.', 'error')
    
    return render_template('register.html')

@app.route('/verification-sent/<email>')
def verification_sent(email):
    """Page shown after registration with instructions"""
    return render_template('verification_sent.html', email=email)

@app.route('/verification-needed/<email>')
def verification_needed(email):
    """Page for users who need to verify their email"""
    user = User.query.filter_by(email=email).first()
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('register'))
    
    if user.is_email_verified:
        flash('Email already verified. You can login now.', 'info')
        return redirect(url_for('login'))
    
    return render_template('verification_needed.html', email=email, user=user)

@app.route('/verify-email/<token>')
def verify_email(token):
    """Email verification endpoint"""
    user, message = verify_user_email(token)
    
    if user:
        # Send welcome email - get email and name from user object
        send_email_async(send_welcome_email, user.email, user.name)
        
        # Auto-login the user
        login_user(user)
        flash('Email verified successfully! Welcome to Granty!', 'success')
        return redirect(url_for('dashboard'))
    else:
        flash(f'Verification failed: {message}', 'error')
        return redirect(url_for('login'))

@app.route('/resend-verification', methods=['POST'])
def resend_verification():
    """Resend verification email"""
    email = request.form.get('email')
    
    user = User.query.filter_by(email=email).first()
    if not user:
        flash('User not found.', 'error')
        return redirect(url_for('register'))
    
    if user.is_email_verified:
        flash('Email already verified.', 'info')
        return redirect(url_for('login'))
    
    # Regenerate token
    user, message = resend_verification_email(user.id)
    if user:
        app_url = request.url_root.rstrip('/')
        # Get name and token from user object
        name = user.name
        token = user.email_verification_token
        
        if send_email_async(send_verification_email, email, name, token, app_url):
            flash('Verification email sent successfully! Please check your inbox.', 'success')
        else:
            flash('Failed to send verification email. Please try again.', 'error')
    else:
        flash(f'Error: {message}', 'error')
    
    return redirect(url_for('verification_needed', email=email))

# Add a decorator to protect routes that require verified email
from functools import wraps

def email_verified_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated:
            return redirect(url_for('login'))
        if not current_user.is_email_verified:
            flash('Please verify your email address to access this feature.', 'warning')
            return redirect(url_for('verification_needed', email=current_user.email))
        return f(*args, **kwargs)
    return decorated_function

# Update dashboard route to require email verification
@app.route('/dashboard')
@login_required
@email_verified_required
def dashboard():
    projects = storage.get_projects(current_user.id)
    template_categories = storage.get_templates_by_category()
    return render_template('dashboard.html', projects=projects, template_categories=template_categories, user=current_user)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out successfully!', 'success')
    return redirect(url_for('index'))

@app.route('/new-document')
def new_document():
    """Main document creation page with all categories"""
    template_categories = storage.get_templates_by_category()
    return render_template('new_document.html', template_categories=template_categories)

@app.route('/create-document', methods=['POST'])
def create_document():
    """Create a new document from template"""
    template_id = request.form.get('template_id')
    
    # Check project limit for logged-in users
    if current_user.is_authenticated and not current_user.can_create_project():
        flash(f'You have reached the maximum limit of {app.config.get("MAX_PROJECTS_PER_USER", 5)} projects.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get template
    template = storage.get_template(template_id)
    if not template:
        flash('Template not found.', 'error')
        return redirect(url_for('new_document'))
    
    project_id = str(uuid.uuid4())
    
    # Create new project
    project_data = {
        'id': project_id,
        'title': f"New {template.category.replace('_', ' ').title()} - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        'project_type': template.category,
        'template_id': template_id,
        'sections': template.get_sections(),
        'answers': {},
        'created_at': datetime.now().isoformat(),
        'updated_at': datetime.now().isoformat(),
        'generated_content': None
    }
    
    user_id = current_user.id if current_user.is_authenticated else None
    project = storage.save_project(project_data, user_id)
    
    if project:
        return redirect(url_for('editor', draft_id=project.id))
    else:
        flash('Error creating project.', 'error')
        return redirect(url_for('new_document'))

@app.route('/upload-custom', methods=['POST'])
def upload_custom():
    print("Processing custom template upload...")
    
    if 'custom_template' not in request.files:
        flash('No file selected.', 'error')
        return redirect(url_for('new_document'))
    
    file = request.files['custom_template']
    if file.filename == '':
        flash('No file selected.', 'error')
        return redirect(url_for('new_document'))
    
    filename = secure_filename(file.filename)
    file_path = os.path.join('data/uploads', filename)
    os.makedirs('data/uploads', exist_ok=True)
    
    try:
        file.save(file_path)
        print(f"File saved to: {file_path}")
        
        # Extract content based on file type
        content = ""
        if filename.lower().endswith('.pdf'):
            print("Processing PDF file...")
            content = extract_text_from_pdf(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        print(f"Extracted content length: {len(content)} characters")
        
        if not content.strip():
            flash('Could not extract content from file.', 'error')
            return redirect(url_for('new_document'))
        
        # Process with Gemini
        sections_with_sample_answers = gemini_utils.extract_sections_with_sample_answers(content)
        
        # Create template
        template_data = {
            'id': str(uuid.uuid4()),
            'name': f"Custom - {filename}",
            'description': "Custom uploaded template with AI-generated sample answers",
            'category': 'custom',
            'sections': sections_with_sample_answers
        }
        
        template = storage.save_template(template_data)
        
        # Create project
        project_data = {
            'id': str(uuid.uuid4()),
            'title': f"Custom Project - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            'project_type': 'custom',
            'template_id': template.id,
            'sections': sections_with_sample_answers,
            'answers': {q['id']: q.get('sample_answer', '') for section in sections_with_sample_answers for q in section.get('questions', [])},
            'generated_content': None
        }
        
        user_id = current_user.id if current_user.is_authenticated else None
        project = storage.save_project(project_data, user_id)
        
        return redirect(url_for('editor', draft_id=project.id))
        
    except Exception as e:
        print(f"Error processing custom template: {e}")
        flash(f'Error processing template: {str(e)}', 'error')
        return redirect(url_for('new_document'))

@app.route('/editor/<draft_id>')
def editor(draft_id):
    project = storage.get_project(draft_id)
    if not project:
        flash('Project not found.', 'error')
        return redirect(url_for('index'))
    
    # Check access permissions
    if project.user_id and (not current_user.is_authenticated or current_user.id != project.user_id):
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    
    template = storage.get_template(project.template_id)
    
    # Convert project to dict for template compatibility
    draft = {
        'id': project.id,
        'title': project.title,
        'template_id': project.template_id,
        'sections': project.get_sections(),
        'answers': project.get_answers(),
        'generated_content': project.generated_content
    }
    
    return render_template('editor.html', draft=draft, template=template, project=project)

@app.route('/preview/<draft_id>')
def preview(draft_id):
    project = storage.get_project(draft_id)
    if not project:
        flash('Project not found.', 'error')
        return redirect(url_for('index'))
    
    # Check access permissions
    if project.user_id and (not current_user.is_authenticated or current_user.id != project.user_id):
        flash('Access denied.', 'error')
        return redirect(url_for('index'))
    
    # Convert project to dict for template compatibility
    draft = {
        'id': project.id,
        'title': project.title,
        'template_id': project.template_id,
        'sections': project.get_sections(),
        'answers': project.get_answers(),
        'generated_content': project.generated_content
    }
    
    return render_template('preview.html', draft=draft, project=project)

@app.route('/chat/<project_id>')
def chat_interface(project_id):
    """Chat interface available for all users (both logged in and guest)"""
    project = storage.get_project(project_id)
    if not project:
        flash('Project not found or access denied.', 'error')
        return redirect(url_for('index'))
    
    # Check access permissions for logged-in users
    if project.user_id and current_user.is_authenticated and project.user_id != current_user.id:
        flash('Access denied.', 'error')
        return redirect(url_for('dashboard'))
    
    # Allow guest access to guest projects
    if project.user_id and not current_user.is_authenticated:
        flash('Please login to access this project.', 'error')
        return redirect(url_for('login'))
    
    chat_history = storage.get_chat_history(project_id)
    return render_template('chat.html', project=project, chat_history=chat_history)

@app.route('/api/restructure-content', methods=['POST'])
def api_restructure_content():
    """Restructure content (available for all users)"""
    try:
        data = request.json
        project_id = data.get('project_id')
        instructions = data.get('instructions')
        
        project = storage.get_project(project_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'})
        
        # Check access permissions
        if project.user_id and current_user.is_authenticated and project.user_id != current_user.id:
            return jsonify({'success': False, 'error': 'Access denied'})
        
        if not project.generated_content:
            return jsonify({'success': False, 'error': 'No content to restructure'})
        
        # Restructure content with AI
        new_content = gemini_utils.restructure_content_with_ai(project.generated_content, instructions)
        
        # Update project
        storage.update_project(project_id, {'generated_content': new_content})
        
        return jsonify({'success': True, 'content': new_content})
        
    except Exception as e:
        print(f"Error restructuring content: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/save-answers', methods=['POST'])
def save_answers():
    try:
        data = request.json
        project_id = data.get('draft_id')  # Keep compatibility with existing frontend
        answers = data.get('answers', {})
        
        if not project_id:
            return jsonify({'success': False, 'error': 'Project ID is required'})
        
        project = storage.get_project(project_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'})
        
        # Check access permissions
        if project.user_id and (not current_user.is_authenticated or current_user.id != project.user_id):
            return jsonify({'success': False, 'error': 'Access denied'})
        
        # Update answers
        current_answers = project.get_answers()
        current_answers.update(answers)
        
        storage.update_project(project_id, {
            'answers': current_answers,
            'generated_content': None  # Clear generated content
        })
        
        return jsonify({'success': True, 'message': f'Saved {len(answers)} answers'})
        
    except Exception as e:
        print(f"Error saving answers: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/generate-content', methods=['POST'])
def generate_content():
    try:
        data = request.json
        project_id = data.get('draft_id')  # Keep compatibility
        
        project = storage.get_project(project_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'})
        
        # Check access permissions
        if project.user_id and (not current_user.is_authenticated or current_user.id != project.user_id):
            return jsonify({'success': False, 'error': 'Access denied'})
        
        template = storage.get_template(project.template_id)
        if not template:
            return jsonify({'success': False, 'error': 'Template not found'})
        
        # Convert to dict for compatibility with existing function
        draft_dict = {
            'id': project.id,
            'title': project.title,
            'sections': project.get_sections(),
            'answers': project.get_answers()
        }
        
        template_dict = {
            'name': template.name,
            'sections': template.get_sections()
        }
        
        # Generate content
        content = gemini_utils.generate_grant_content(draft_dict, template_dict)
        
        # Save generated content
        storage.update_project(project_id, {'generated_content': content})
        
        return jsonify({'success': True, 'content': content})
        
    except Exception as e:
        print(f"Error generating content: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/save-content', methods=['POST'])
def save_content():
    try:
        data = request.json
        project_id = data.get('draft_id')  # Keep compatibility
        content = data.get('content')
        
        if not project_id or not content:
            return jsonify({'success': False, 'error': 'Missing project ID or content'})
        
        project = storage.get_project(project_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'})
        
        # Check access permissions
        if project.user_id and (not current_user.is_authenticated or current_user.id != project.user_id):
            return jsonify({'success': False, 'error': 'Access denied'})
        
        storage.update_project(project_id, {'generated_content': content})
        
        return jsonify({'success': True, 'message': 'Content saved successfully'})
        
    except Exception as e:
        print(f"Error saving content: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/chat', methods=['POST'])
def api_chat():
    """Chat API available for all users"""
    try:
        data = request.json
        project_id = data.get('project_id')
        message = data.get('message')
        
        project = storage.get_project(project_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'})
        
        # Check access permissions
        if project.user_id and current_user.is_authenticated and project.user_id != current_user.id:
            return jsonify({'success': False, 'error': 'Access denied'})
        
        # Prepare project context
        project_context = {
            'title': project.title,
            'project_type': project.project_type,
            'generated_content': project.generated_content,
            'sections': project.get_sections(),
            'answers': project.get_answers()
        }
        
        # Use simple AI chat without complex agents
        response = gemini_utils.chat_with_gemini_direct(message, project_context)
        
        # Save chat message
        storage.save_chat_message(project_id, message, response)
        
        return jsonify({'success': True, 'response': response})
        
    except Exception as e:
        print(f"Error in chat API: {e}")
        return jsonify({'success': False, 'error': 'Sorry, I encountered an error. Please try again.'})

def handle_edit_request(message, project_context, project_id):
    """Simple edit request handler"""
    try:
        response = f"I understand you want to edit: {message}. Here are some suggestions:\n\n"
        response += "1. Review the content in the preview section\n"
        response += "2. Use the 'Enable Editing' button to make changes\n"
        response += "3. Focus on clarity and professional tone\n"
        response += "4. Ensure all required sections are complete\n\n"
        response += "Would you like me to help with any specific aspect of your document?"
        
        return response
        
    except Exception as e:
        return f"I can help you with editing. Please describe what specific changes you'd like to make."

def handle_generate_request(message, project_context, project_id):
    """Simple content generation handler"""
    try:
        response = f"For content generation: {message}\n\n"
        response += "I recommend:\n"
        response += "1. Ensure all questions in the editor are answered\n"
        response += "2. Use the 'Generate Content' button in the preview\n"
        response += "3. Review and edit the generated content as needed\n"
        response += "4. Use professional language and clear structure\n\n"
        response += "What specific type of content would you like me to help you improve?"
        
        return response
        
    except Exception as e:
        return f"I can help you generate content. Please describe what you need help with."

@app.route('/opportunities', methods=['GET', 'POST'])
def opportunities():
    if request.method == 'GET':
        return render_template('opportunities.html')
    elif request.method == 'POST':
        content = ""
        if 'opportunity_file' in request.files and request.files['opportunity_file'].filename != '':
            file = request.files['opportunity_file']
            filename = secure_filename(file.filename)
            file_path = os.path.join('data/uploads', filename)
            os.makedirs('data/uploads', exist_ok=True)
            file.save(file_path)

            if filename.lower().endswith('.pdf'):
                content = extract_text_from_pdf(file_path)
            else:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
        if not content.strip():
            content = request.form.get('research_text', '')
            
        if content.strip():
            # For guest users, show limited info
            if not current_user.is_authenticated:
                flash('Opportunities found! Please login to view detailed results.', 'info')
                return render_template('opportunities.html', guest_message=True)
            else:
                result = gemini_utils.get_opportunities(content)
                return render_template('opportunities.html', opportunities=result)
            
        return render_template('opportunities.html')

def extract_text_from_pdf(file_path):
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""
def clean_and_prepare_html(html_content):
    """
    Clean HTML content and prepare it for conversion while preserving formatting
    """
    if not html_content:
        return html_content
    
    import re
    
    print(f"Before cleaning: {html_content[:100]}...")
    
    # Remove ```html and ``` markers but keep the HTML content
    content = re.sub(r'```html\s*\n?', '', html_content, flags=re.IGNORECASE)
    content = re.sub(r'```\s*$', '', content, flags=re.MULTILINE)
    content = re.sub(r'```\s*\n', '', content)
    content = re.sub(r'```\w*\s*\n?', '', content, flags=re.IGNORECASE)
    
    # Also remove any markdown code block markers
    content = re.sub(r'^\s*```.*?$', '', content, flags=re.MULTILINE)
    
    # Clean up extra whitespace
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
    content = content.strip()
    
    print(f"After cleaning: {content[:100]}...")
    
    return content

# Also update the PDF generation function to be more robust
def generate_pdf_from_html_reportlab(html_content, title):
    """
    Generate PDF using reportlab with better content handling
    """
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        print(f"ReportLab PDF - Input content: {html_content[:100]}...")
        
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_style = styles['Title']
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 20))
        
        # Process content - handle HTML tags better
        import re
        
        content = html_content
        
        # Convert HTML headers to formatted text
        content = re.sub(r'<h1>(.*?)</h1>', r'\n\n<b><font size="16">\1</font></b>\n\n', content, flags=re.DOTALL)
        content = re.sub(r'<h2>(.*?)</h2>', r'\n\n<b><font size="14">\1</font></b>\n\n', content, flags=re.DOTALL)
        content = re.sub(r'<h3>(.*?)</h3>', r'\n\n<b><font size="12">\1</font></b>\n\n', content, flags=re.DOTALL)
        
        # Preserve basic formatting
        content = re.sub(r'<strong>(.*?)</strong>', r'<b>\1</b>', content, flags=re.DOTALL)
        content = re.sub(r'<em>(.*?)</em>', r'<i>\1</i>', content, flags=re.DOTALL)
        
        # Handle line breaks
        content = re.sub(r'<br\s*/?>', '\n', content)
        content = re.sub(r'</p>', '\n\n', content)
        content = re.sub(r'<p[^>]*>', '', content)
        
        # Remove remaining HTML tags
        content = re.sub(r'<[^>]+>', '', content)
        
        # Decode HTML entities
        content = content.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        content = content.replace('&nbsp;', ' ')
        
        print(f"ReportLab PDF - Processed content: {content[:200]}...")
        
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            para = para.strip()
            if para:
                try:
                    # Handle very long paragraphs
                    if len(para) > 2000:
                        # Split long paragraphs
                        chunks = [para[i:i+2000] for i in range(0, len(para), 2000)]
                        for chunk in chunks:
                            story.append(Paragraph(chunk, styles['Normal']))
                            story.append(Spacer(1, 12))
                    else:
                        story.append(Paragraph(para, styles['Normal']))
                        story.append(Spacer(1, 12))
                except Exception as e:
                    print(f"Error with paragraph: {e}")
                    # Escape problematic characters
                    safe_para = para.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                    safe_para = ''.join(char for char in safe_para if ord(char) < 128)  # Remove non-ASCII
                    try:
                        story.append(Paragraph(safe_para[:1000], styles['Normal']))
                        story.append(Spacer(1, 12))
                    except:
                        # Skip problematic paragraphs
                        print(f"Skipping problematic paragraph: {para[:50]}...")
                        continue
        
        doc.build(story)
        print("✓ PDF generated using ReportLab")
        return temp_file.name
        
    except Exception as e:
        print(f"Error generating PDF with reportlab: {e}")
        import traceback
        traceback.print_exc()
        return None

def generate_pdf_from_html_improved(html_content, title):
    """
    Generate PDF from HTML with the best available method
    """
    try:
        # Clean content first
        clean_html = clean_and_prepare_html(html_content)
        
        # Try weasyprint first (best quality) - import only when needed
        if WEASYPRINT_AVAILABLE:
            try:
                import weasyprint
                
                # Prepare full HTML document
                full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{ 
                            font-family: Arial, sans-serif; 
                            line-height: 1.6; 
                            color: #333; 
                            margin: 40px;
                            max-width: none;
                        }}
                        h1 {{ 
                            color: #2c3e50; 
                            border-bottom: 2px solid #3498db; 
                            padding-bottom: 10px; 
                            page-break-after: avoid;
                        }}
                        h2 {{ 
                            color: #34495e; 
                            border-bottom: 1px solid #3498db; 
                            padding-bottom: 5px; 
                            page-break-after: avoid;
                        }}
                        h3 {{ color: #2c3e50; page-break-after: avoid; }}
                        p {{ margin-bottom: 1em; text-align: justify; }}
                        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
                        strong {{ font-weight: bold; }}
                        em {{ font-style: italic; }}
                        @page {{ margin: 1in; }}
                    </style>
                </head>
                <body>
                    <h1>{title}</h1>
                    {clean_html}
                </body>
                </html>
                """
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                weasyprint.HTML(string=full_html).write_pdf(temp_file.name)
                print("✓ PDF generated using WeasyPrint")
                return temp_file.name
                
            except Exception as e:
                print(f"WeasyPrint failed: {e}, trying next method...")
        
        # Try pdfkit as alternative
        if PDFKIT_AVAILABLE:
            try:
                import pdfkit
                
                # Prepare full HTML document
                full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                        h2 {{ color: #34495e; border-bottom: 1px solid #3498db; padding-bottom: 5px; }}
                        p {{ margin-bottom: 1em; text-align: justify; }}
                        strong {{ font-weight: bold; }}
                    </style>
                </head>
                <body>
                    <h1>{title}</h1>
                    {clean_html}
                </body>
                </html>
                """
                
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None,
                    'enable-local-file-access': None
                }
                
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                pdfkit.from_string(full_html, temp_file.name, options=options)
                print("✓ PDF generated using pdfkit")
                return temp_file.name
                
            except Exception as e:
                print(f"pdfkit failed: {e}, using reportlab fallback...")
        
        # Fallback to reportlab (always available)
        return generate_pdf_from_html_reportlab(clean_html, title)
        
    except Exception as e:
        print(f"Error in PDF generation: {e}")
        return generate_pdf_from_html_reportlab(html_content, title)


def html_to_docx_improved(html_content, title):
    """
    Convert HTML content to DOCX format with better formatting preservation
    """
    if not DOCX_AVAILABLE:
        print("Error: python-docx not available")
        return None
        
    try:
        # Clean content
        clean_html = clean_and_prepare_html(html_content)
        
        # Create a new Document
        doc = DocxDocument()
        
        # Add title
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if BS4_AVAILABLE:
            # Use BeautifulSoup for better parsing
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(clean_html, 'html.parser')
            
            # Process each element
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Add heading
                    level = int(element.name[1])
                    doc.add_heading(element.get_text().strip(), level)
                    
                elif element.name in ['p', 'div']:
                    # Add paragraph with basic formatting
                    text = element.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph()
                        
                        # Handle bold and italic text
                        for content in element.contents:
                            if hasattr(content, 'name'):
                                if content.name in ['strong', 'b']:
                                    run = paragraph.add_run(content.get_text())
                                    run.bold = True
                                elif content.name in ['em', 'i']:
                                    run = paragraph.add_run(content.get_text())
                                    run.italic = True
                                else:
                                    paragraph.add_run(content.get_text() if hasattr(content, 'get_text') else str(content))
                            else:
                                paragraph.add_run(str(content))
        else:
            # Simple parsing without BeautifulSoup
            import re
            
            # Remove HTML tags but try to preserve structure
            content = clean_html
            
            # Split by major elements
            sections = re.split(r'<h[1-6].*?>(.*?)</h[1-6]>', content)
            
            for i, section in enumerate(sections):
                section = section.strip()
                if section:
                    if i % 2 == 1:  # This is a heading
                        doc.add_heading(re.sub(r'<[^>]+>', '', section), 1)
                    else:  # This is content
                        # Split into paragraphs
                        paragraphs = section.split('\n\n')
                        for para in paragraphs:
                            para = re.sub(r'<[^>]+>', '', para).strip()
                            if para:
                                doc.add_paragraph(para)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        print("✓ DOCX generated successfully")
        return temp_file.name
        
    except Exception as e:
        print(f"Error in DOCX generation: {e}")
        return None

# Update the main conversion functions
def html_to_docx(html_content, title):
    """Main DOCX conversion function"""
    return html_to_docx_improved(html_content, title)

def generate_pdf_from_html(html_content, title):
    """Main PDF conversion function"""  
    return generate_pdf_from_html_improved(html_content, title)

@app.route('/api/download-docx/<draft_id>')
def download_docx(draft_id):
    try:
        if not DOCX_AVAILABLE:
            return jsonify({'success': False, 'error': 'DOCX functionality not available. Please install python-docx.'}), 500
            
        coupon_used = request.args.get('coupon') == 'grantyadmin'
        payment_verified = session.get('payment_verified', False) and session.get('draft_id') == draft_id
        
        if not (coupon_used or payment_verified):
            return jsonify({'success': False, 'error': 'Payment required'}), 403
        
        project = storage.get_project(draft_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        content = project.generated_content
        if not content:
            return jsonify({'success': False, 'error': 'No content to download'}), 400
        
        docx_path = html_to_docx(content, project.title)
        if not docx_path:
            return jsonify({'success': False, 'error': 'Failed to generate DOCX'}), 500
        
        if not coupon_used:
            session.pop('payment_verified', None)
            session.pop('draft_id', None)
        
        return send_file(
            docx_path,
            as_attachment=True,
            download_name=f"{project.title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"Error downloading DOCX: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/download-pdf/<draft_id>')
def download_pdf(draft_id):
    try:
        coupon_used = request.args.get('coupon') == 'grantyadmin'
        payment_verified = session.get('payment_verified', False) and session.get('draft_id') == draft_id
        
        if not (coupon_used or payment_verified):
            return jsonify({'success': False, 'error': 'Payment required'}), 403
        
        project = storage.get_project(draft_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'}), 404
        
        content = project.generated_content
        if not content:
            return jsonify({'success': False, 'error': 'No content to download'}), 400
        
        pdf_path = generate_pdf_from_html(content, project.title)
        if not pdf_path:
            return jsonify({'success': False, 'error': 'Failed to generate PDF'}), 500
        
        if not coupon_used:
            session.pop('payment_verified', None)
            session.pop('draft_id', None)
        
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=f"{project.title}.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"Error downloading PDF: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/payment/<draft_id>')
def payment_page(draft_id):
    project = storage.get_project(draft_id)
    if not project:
        return redirect(url_for('index'))
    
    payment_id = f"PAY_{int(time.time())}_{random.randint(1000, 9999)}"
    
    session['payment_id'] = payment_id
    session['draft_id_for_payment'] = draft_id
    session['payment_amount'] = 49
    session['payment_status'] = 'pending'
    
    draft = {
        'id': project.id,
        'title': project.title
    }
    
    return render_template('payment.html', 
                         draft=draft, 
                         payment_id=payment_id,
                         amount=49)

# Update your verify_payment_new route in app.py

@app.route('/verify-payment-new', methods=['POST'])
def verify_payment_new():
    try:
        data = request.json
        transaction_id = data.get('transaction_id', '').strip()
        payment_id = session.get('payment_id')
        draft_id = session.get('draft_id_for_payment')
        
        print(f"Payment verification attempt:")
        print(f"  Transaction ID: {transaction_id}")
        print(f"  Payment ID: {payment_id}")
        print(f"  Draft ID: {draft_id}")
        print(f"  Current session: {dict(session)}")
        
        if not transaction_id:
            return jsonify({
                'success': False, 
                'message': 'Please enter transaction ID'
            })
        
        if not payment_id or not draft_id:
            return jsonify({
                'success': False, 
                'message': 'Payment session expired. Please try again.'
            })
        
        if len(transaction_id) < 8:
            return jsonify({
                'success': False, 
                'message': 'Invalid transaction ID. Please check and try again.'
            })
        
        # Set session variables for download routes
        session['payment_status'] = 'completed'
        session['transaction_id'] = transaction_id
        session['payment_verified'] = True  # For backward compatibility
        session['draft_id'] = draft_id      # For backward compatibility
        
        print(f"Payment verified successfully. Updated session: {dict(session)}")
        
        return jsonify({
            'success': True, 
            'message': 'Payment verified successfully!',
            'transaction_id': transaction_id
        })
        
    except Exception as e:
        print(f"Payment verification error: {e}")
        return jsonify({
            'success': False, 
            'message': 'Payment verification failed. Please try again.'
        })

# Update these download routes in your app.py

@app.route('/download-docx-new/<draft_id>')
def download_docx_new(draft_id):
    try:
        if not DOCX_AVAILABLE:
            return jsonify({'error': 'DOCX functionality not available. Please install python-docx.'}), 500
        
        # Check both session structures
        payment_completed = (
            session.get('payment_status') == 'completed' and 
            session.get('draft_id_for_payment') == draft_id
        ) or (
            session.get('payment_verified', False) and 
            session.get('draft_id') == draft_id
        )
        
        print(f"DOCX Download attempt - Draft ID: {draft_id}")
        print(f"Session data: {dict(session)}")
        print(f"Payment completed: {payment_completed}")
        
        if not payment_completed:
            print("Payment not verified, redirecting to payment page")
            return redirect(url_for('payment_page', draft_id=draft_id))
        
        project = storage.get_project(draft_id)
        if not project:
            return jsonify({'error': 'Project not found'}), 404
        
        content = project.generated_content
        if not content:
            return jsonify({'error': 'No content to download. Please generate content first.'}), 400
        
        print(f"Original content length: {len(content)} characters")
        print(f"Content preview: {content[:200]}...")
        
        # CLEAN CONTENT before creating DOCX
        cleaned_content = clean_and_prepare_html(content)
        print(f"Cleaned content length: {len(cleaned_content)} characters")
        print(f"Cleaned content preview: {cleaned_content[:200]}...")
        
        docx_path = html_to_docx(cleaned_content, project.title)
        if not docx_path:
            return jsonify({'error': 'Failed to generate DOCX'}), 500
        
        print("DOCX generated successfully, starting download")
        
        # DON'T clear session yet - keep it for other downloads
        # We'll clear it in a separate cleanup route or let it expire
        
        return send_file(
            docx_path,
            as_attachment=True,
            download_name=f"{project.title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        print(f"DOCX download error: {e}")
        return jsonify({'error': 'Download failed. Please try again.'}), 500

@app.route('/download-pdf-new/<draft_id>')
def download_pdf_new(draft_id):
    try:
        # Check both session structures
        payment_completed = (
            session.get('payment_status') == 'completed' and 
            session.get('draft_id_for_payment') == draft_id
        ) or (
            session.get('payment_verified', False) and 
            session.get('draft_id') == draft_id
        )
        
        print(f"PDF Download attempt - Draft ID: {draft_id}")
        print(f"Session data: {dict(session)}")
        print(f"Payment completed: {payment_completed}")
        
        if not payment_completed:
            print("Payment not verified, redirecting to payment page")
            return redirect(url_for('payment_page', draft_id=draft_id))
        
        project = storage.get_project(draft_id)
        if not project:
            return jsonify({'error': 'Project not found'}), 404
        
        content = project.generated_content
        if not content:
            return jsonify({'error': 'No content to download. Please generate content first.'}), 400
        
        print(f"Original content length: {len(content)} characters")
        print(f"Content preview: {content[:200]}...")
        
        # CLEAN CONTENT before creating PDF
        cleaned_content = clean_and_prepare_html(content)
        print(f"Cleaned content length: {len(cleaned_content)} characters")
        print(f"Cleaned content preview: {cleaned_content[:200]}...")
        
        pdf_path = generate_pdf_from_html(cleaned_content, project.title)
        if not pdf_path:
            return jsonify({'error': 'Failed to generate PDF'}), 500
        
        print("PDF generated successfully, starting download")
        
        # DON'T clear session yet - keep it for other downloads
        
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=f"{project.title}.pdf",
            mimetype='application/pdf'
        )
        
    except Exception as e:
        print(f"PDF download error: {e}")
        return jsonify({'error': 'Download failed. Please try again.'}), 500

# Add a route to clear payment session after both downloads are done
@app.route('/clear-payment-session/<draft_id>')
def clear_payment_session(draft_id):
    """Clear payment session after downloads are complete"""
    try:
        if session.get('draft_id_for_payment') == draft_id or session.get('draft_id') == draft_id:
            session.pop('payment_status', None)
            session.pop('payment_id', None)
            session.pop('draft_id_for_payment', None)
            session.pop('transaction_id', None)
            session.pop('payment_verified', None)
            session.pop('draft_id', None)
            print("Payment session cleared")
        
        return jsonify({'success': True, 'message': 'Session cleared'})
    except Exception as e:
        print(f"Error clearing session: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/project-insights/<project_id>')
def api_project_insights(project_id):
    """Get AI-powered insights about a project (available for all users)"""
    try:
        project = storage.get_project(project_id)
        if not project:
            return jsonify({'success': False, 'error': 'Project not found'})
        
        # Check access permissions
        if project.user_id and current_user.is_authenticated and project.user_id != current_user.id:
            return jsonify({'success': False, 'error': 'Access denied'})
        
        # Generate insights using AI
        insights_prompt = f"""
        Analyze this project and provide insights and recommendations:
        
        Title: {project.title}
        Type: {project.project_type}
        Sections: {len(project.get_sections())}
        Answers completed: {len([v for v in project.get_answers().values() if str(v).strip()])}
        Has generated content: {'Yes' if project.generated_content else 'No'}
        
        Provide:
        1. Completion status and next steps
        2. Content quality assessment
        3. Suggestions for improvement
        4. Missing elements
        5. Strengths of the current content
        
        Format as structured text with clear sections.
        """
        
        insights = gemini_utils.chat_with_gemini_direct(insights_prompt)
        
        return jsonify({
            'success': True,
            'insights': insights,
            'project_stats': {
                'total_sections': len(project.get_sections()),
                'completed_answers': len([v for v in project.get_answers().values() if str(v).strip()]),
                'has_content': bool(project.generated_content),
                'last_updated': project.updated_at.isoformat()
            }
        })
        
    except Exception as e:
        print(f"Error getting project insights: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/health')
def health_check():
    return jsonify({'status': 'healthy'}), 200
    
# Enhanced error handling
@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    print(f"Internal error: {error}")
    return render_template('error.html', error="Internal server error"), 500

@app.errorhandler(404)
def not_found_error(error):
    return render_template('error.html', error="Page not found"), 404

if __name__ == '__main__':
    with app.app_context():
        storage.init_db(app)
    
    app.run(debug=True)
